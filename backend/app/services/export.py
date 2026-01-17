# from pathlib import Path
# from uuid import uuid4
# from docx import Document
# from reportlab.platypus import SimpleDocTemplate, Paragraph
# from reportlab.lib.styles import getSampleStyleSheet
# from reportlab.lib.pagesizes import A4

# TEMP_DIR = Path("temp")
# TEMP_DIR.mkdir(exist_ok=True)

# def export_docx(text: str) -> Path:
#     path = TEMP_DIR / f"{uuid4().hex}.docx"
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save(path)
#     return path

# def export_pdf(text: str) -> Path:
#     path = TEMP_DIR / f"{uuid4().hex}.pdf"
#     styles = getSampleStyleSheet()
#     doc = SimpleDocTemplate(str(path), pagesize=A4)
#     doc.build([Paragraph(text, styles["Normal"])])
#     return path


from pathlib import Path
from uuid import uuid4
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)

def export_docx(text: str) -> Path:
    path = EXPORT_DIR / f"{uuid4().hex}.docx"
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
    return path

def export_pdf_from_markdown(text: str) -> Path:
    path = EXPORT_DIR / f"{uuid4().hex}.pdf"
    styles = getSampleStyleSheet()
    story = []

    for line in text.split("\n"):
        if line.startswith("#"):
            story.append(Paragraph(f"<b>{line.replace('#','')}</b>", styles["Heading2"]))
        else:
            story.append(Paragraph(line, styles["Normal"]))

    SimpleDocTemplate(str(path), pagesize=A4).build(story)
    return path
